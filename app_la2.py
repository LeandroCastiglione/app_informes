import streamlit as st
import pandas as pd
from PIL import Image
from docx import Document
import os

# Función para guardar los datos en un archivo Excel
def save_to_excel(data, filename="auditoria.xlsx"):
    df = pd.DataFrame(data)
    df.to_excel(filename, index=False)

# Función para generar un informe en Word
def generate_report(data, filename="informe_auditoria.docx"):
    doc = Document()
    doc.add_heading("Informe de Auditoría", level=1)
    
    for i, item in enumerate(data):
        doc.add_heading(f"Hallazgo {i+1}", level=2)
        doc.add_paragraph(f"Descripción: {item['descripcion']}")
        if item['foto']:
            doc.add_picture(item['foto'], width=doc.sections[0].page_width - 200)
    
    doc.save(filename)
    return filename

# Configuración de la app
st.title("Relevamiento de Auditoría")

# Lista para almacenar hallazgos
data = []

# Formulario para ingresar datos
descripcion = st.text_area("Descripción del hallazgo")
foto = st.file_uploader("Subir una foto", type=["jpg", "png", "jpeg"])

if st.button("Guardar hallazgo"):
    foto_path = None
    if foto:
        foto_path = f"temp_{foto.name}"
        with open(foto_path, "wb") as f:
            f.write(foto.getbuffer())
    
    data.append({"descripcion": descripcion, "foto": foto_path})
    st.success("Hallazgo guardado!")

# Botón para generar informe
if st.button("Generar informe"):
    if data:
        report_path = generate_report(data)
        st.download_button(
            label="Descargar Informe", 
            data=open(report_path, "rb"),
            file_name="informe_auditoria.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("No hay datos para generar el informe.")

