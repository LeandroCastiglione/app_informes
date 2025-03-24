import streamlit as st
from docx import Document
from PIL import Image
import io

# Título de la app
st.title("Relevamiento de Hallazgos")

# Inicializar lista de hallazgos en session_state
if "hallazgos" not in st.session_state:
    st.session_state.hallazgos = []

# Formulario para cargar hallazgos
st.header("Agregar un nuevo hallazgo")
descripcion = st.text_area("Descripción del hallazgo")
foto = st.file_uploader("Subir foto", type=["png", "jpg", "jpeg"])

if st.button("Agregar Hallazgo"):
    if descripcion:
        nuevo_hallazgo = {"descripcion": descripcion, "foto": foto}
        st.session_state.hallazgos.append(nuevo_hallazgo)
        st.success("Hallazgo agregado correctamente.")
    else:
        st.warning("Por favor, ingresa una descripción.")

# Mostrar hallazgos cargados
st.header("Hallazgos registrados")
if st.session_state.hallazgos:
    for idx, hallazgo in enumerate(st.session_state.hallazgos):
        st.subheader(f"Hallazgo {idx+1}")
        st.write(hallazgo["descripcion"])
        if hallazgo["foto"]:
            image = Image.open(hallazgo["foto"])
            st.image(image, caption="Foto adjunta", use_column_width=True)

# Función para generar el informe
def generar_informe():
    if not st.session_state.hallazgos:
        st.error("No hay datos para generar el informe.")
        return None

    doc = Document()
    doc.add_heading("Informe de Hallazgos", level=1)

    for idx, hallazgo in enumerate(st.session_state.hallazgos):
        doc.add_heading(f"Hallazgo {idx+1}", level=2)
        doc.add_paragraph(hallazgo["descripcion"])

        if hallazgo["foto"]:
            image = Image.open(hallazgo["foto"])
            img_stream = io.BytesIO()
            image.save(img_stream, format="PNG")
            doc.add_picture(img_stream, width=doc.sections[0].page_width * 0.5)

    # Guardar el archivo en memoria y permitir descarga
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Botón para generar el informe
st.header("Generar Informe")
if st.button("Descargar Informe"):
    informe = generar_informe()
    if informe:
        st.download_button(
            label="📄 Descargar Informe",
            data=informe,
            file_name="informe_hallazgos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Botón para limpiar hallazgos si es necesario
if st.button("Limpiar Hallazgos"):
    st.session_state.hallazgos = []
    st.success("Los hallazgos han sido eliminados.")
# comentario para git a ver si se actualiza.